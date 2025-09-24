# backend/app/core/document_processor.py

from typing import List
import re
import os

class DocumentProcessor:
    """Handles processing of different document types"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_document(self, file) -> List[str]:
        """Process uploaded document and return text chunks"""
        try:
            # Determine file type and extract text
            if hasattr(file, 'type'):
                file_type = file.type
            else:
                # Fallback to extension-based detection
                ext = os.path.splitext(file.name)[1].lower()
                type_map = {'.pdf': 'application/pdf', '.txt': 'text/plain', '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}
                file_type = type_map.get(ext, 'unknown')
            
            if file_type == "application/pdf":
                text = self._extract_pdf_text(file)
            elif file_type == "text/plain":
                text = self._extract_txt_text(file)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = self._extract_docx_text(file)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Clean and chunk the text
            cleaned_text = self._clean_text(text)
            chunks = self._create_chunks(cleaned_text, file.name)
            
            return chunks
            
        except Exception as e:
            print(f"Error processing {file.name}: {str(e)}")
            return []
    
    def _extract_pdf_text(self, file) -> str:
        import PyPDF2
        """Extract text from PDF file"""
        try:
            if hasattr(file, 'read'):
                # Handle file-like object
                pdf_reader = PyPDF2.PdfReader(file)
            else:
                # Handle file path
                with open(file.path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
            
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_txt_text(self, file) -> str:
        """Extract text from TXT file"""
        try:
            if hasattr(file, 'read'):
                # Handle file-like object
                content = file.read()
                if isinstance(content, bytes):
                    # Try different encodings
                    encodings = ['utf-8', 'latin-1', 'cp1252']
                    for encoding in encodings:
                        try:
                            return content.decode(encoding)
                        except UnicodeDecodeError:
                            continue
                    raise Exception("Could not decode file with any supported encoding")
                else:
                    return content
            else:
                # Handle file path
                with open(file.path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def _extract_docx_text(self, file) -> str:
        import docx
        """Extract text from DOCX file"""
        try:
            if hasattr(file, 'read'):
                doc = docx.Document(file)
            else:
                doc = docx.Document(file.path)
            
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s.,!?;:()\-"]', ' ', text)
        text = re.sub(r' +', ' ', text)
        return text.strip()
    
    def _create_chunks(self, text: str, source_name: str) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end < len(text):
                sentence_end = text.rfind('.', start, end)
                if sentence_end != -1 and sentence_end > start + self.chunk_size // 2:
                    end = sentence_end + 1
                else:
                    word_end = text.rfind(' ', start, end)
                    if word_end != -1 and word_end > start + self.chunk_size // 2:
                        end = word_end
            
            chunk = text[start:end].strip()
            
            if chunk:
                chunk_with_metadata = f"[Source: {source_name}]\n{chunk}"
                chunks.append(chunk_with_metadata)
            
            start = end - self.chunk_overlap
            if start <= 0:
                start = end
        
        return chunks

